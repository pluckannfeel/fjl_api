import os
import boto3
import shutil
from dotenv import load_dotenv
from datetime import datetime
from app.helpers.s3_file_upload import generate_s3_url
from tempfile import TemporaryDirectory
from io import BytesIO
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# from lxml import etree
# from zipfile import ZipFile
import logging
# import openpyxl
import xlwings as xw

load_dotenv()

# Initialize S3 client
s3 = boto3.client('s3')

# Define the directory for static files (relative to the current working directory)
STATIC_DIR = "app/static"

contracts_folder = 'uploads/admin/contracts/'
bucket_name = os.getenv("AWS_STORAGE_BUCKET_NAME")


class JobDetails:
    def __init__(self, id: int, title: str, no_of_workers: int, basic_salary: str):
        self.id = id
        self.title = title
        self.no_of_workers = no_of_workers
        self.basic_salary = basic_salary

# Set up logging
# logging.basicConfig(level=logging.DEBUG)

# Function to replace placeholders in a single paragraph, handling split runs


def replace_text_in_paragraph(paragraph, placeholder, replacement):
    paragraph_text = ''.join(run.text for run in paragraph.runs)
    if placeholder in paragraph_text:
        # logging.debug(f"Replacing placeholder in paragraph: {placeholder} -> {replacement}")
        paragraph_text = paragraph_text.replace(placeholder, str(replacement))

        # Clear the existing runs and re-add the modified text
        for run in paragraph.runs:
            run.text = ''
        paragraph.runs[0].text = paragraph_text

# Function to replace placeholders in the entire document


def replace_placeholder(doc, placeholder, replacement):
    # logging.debug(f"Replacing placeholder: {placeholder} -> {replacement}")
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, placeholder, replacement)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholder(cell, placeholder, replacement)

# Function to fill the template with the replacements


def fill_template(template_path, output_path, replacements):
    doc = Document(template_path)
    for placeholder, replacement in replacements.items():
        replace_placeholder(doc, f"{{{{{placeholder}}}}}", replacement)
    doc.save(output_path)


def fill_application_form(company, agency, details):
    created_at = datetime.now()
    year = created_at.year
    month = created_at.month
    day = created_at.day

    # current date
    today = datetime.now()

    visa_type = details['visa_type']

    # new_document_name = "application_form_test.docx"
    if visa_type == 'psw':
        new_document_name = f"MWO_申込書_{company.name_en}_{today.strftime('%Y%m%d')}.docx"
        original_file = os.path.join(STATIC_DIR, "application_form.docx")
    elif visa_type == 'ssw':
        new_document_name = f"MWO_SSW_申込書_{company.name_en}_{today.strftime('%Y%m%d')}.docx"
        original_file = os.path.join(
            STATIC_DIR, 'ssw', 'application_form.docx')
        # check of original_file exists
    else:
        raise ValueError("Invalid visa type")

    s3_new_document = f"{contracts_folder}{new_document_name}"

    # replacements = {
    #     "company_name_en": "Make You Smile Co., Ltd.",
    #     "rep_name_en": "Kazuhiro Ito",
    #     "rep_position_en": "President",
    #     "address": "#304 YT Bashamichi Bldg, 4-20-2 Kaigandori, Naka-ku, Yokohama shi, Kanagawa Pref.",
    #     "postal_code": "231-0002",
    #     "phone": "050-5894-1192",
    #     "website": "https://www.makeyousmile.jp/",
    #     "rep_email": "kazuhiro110@makeyousmile.jp",
    #     "agency_name": "AQIUM INTERNATIONAL INC.",
    #     "agency_address": "Unit 2A, 3A & 3B, 4K Plaza Bldg., 677 Shaw Blvd., Brgy. Kapitolyo, Pasig City, Philippines",
    # }

    replacements = {
        "company_name_en": company.name_en,
        "rep_name_en": company.rep_name_en,
        "rep_position_en": company.rep_position_en,
        "address": f"{company.building_en}, {company.municipality_town_en}, {company.prefecture_en}",
        "postal_code": company.postal_code,
        "phone": company.phone,
        "email": company.email,
        "website": company.website,
        "rep_email": company.rep_email,
        "rep_phone": company.rep_phone,
        "agency_name": agency.name,
        "agency_address": agency.address,
        "agency_rep_name": agency.rep_name,
        "agency_rep_position": agency.rep_position,
    }

    with TemporaryDirectory() as temp_dir:
        temp_file_path = os.path.join(temp_dir, new_document_name)

        new_contract_buffer = BytesIO()
        fill_template(original_file, temp_file_path, replacements)

        with open(temp_file_path, 'rb') as temp_file:
            new_contract_buffer.write(temp_file.read())

        s3.upload_file(temp_file_path, Bucket=bucket_name, Key=s3_new_document, ExtraArgs={
            'ACL': 'public-read', 'ContentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})

    s3_read_url = generate_s3_url(s3_new_document, 'read')
    return s3_read_url


def fill_manpower_request_template(template_path, output_path, replacements, job_details, total_workers):
    doc = Document(template_path)
    for placeholder, replacement in replacements.items():
        replace_placeholder(doc, f"{{{{{placeholder}}}}}", replacement)

    # Handle the job details insertion
    table = doc.tables[0]
    for detail in job_details:
        row_cells = table.add_row().cells
        row_cells[0].text = detail['title']
        row_cells[1].text = str(detail['no_of_workers'])
        row_cells[2].text = detail['basic_salary']

    # Add the total number of workers row
    total_row_cells = table.add_row().cells
    total_row_cells[0].text = "TOTAL Number of Workers"
    total_row_cells[1].text = str(total_workers)
    total_row_cells[2].text = ""

    doc.save(output_path)


def fill_manpower_request_form(company, agency, details):
    # created_at = datetime.now()
    # year = created_at.year
    # month = created_at.month
    # month_name = created_at.strftime("%b")
    # day = created_at.day

    # retrieve created_date from details and convert to date
    if details['created_date'] is None or details['created_date'] == "":
        created_date = datetime.now()
    else:
        created_date = datetime.strptime(
            details['created_date'], "%Y-%m-%dT%H:%M:%S.%fZ")

    year = created_date.year
    month = created_date.month
    month_name = created_date.strftime("%b")
    day = created_date.day

    # current date
    today = datetime.now()

    # new_document_name = "manpower_request_form_test.docx"
    new_document_name = f"MWO_MANPOWER_REQUEST_{company.name_en}_{today.strftime('%Y%m%d')}.docx"
    original_file = os.path.join(STATIC_DIR, "manpower_request.docx")
    s3_new_document = f"{contracts_folder}{new_document_name}"

    # replacements = {
    #     "created_date": f"{day} {month_name}, {year}",
    #     "company_name_en": "Make You Smile Co., Ltd.",
    #     "rep_name_en": "Kazuhiro Ito",
    #     "rep_position_en": "President",
    #     "building": "#304 YT Bashamichi Bldg",
    #     "street": "4-20-2 Kaigandori",
    #     "city": "Naka-ku, Yokohama shi",
    #     "prefecture": "Kanagawa Pref.",
    #     "postal_code": "231-0002",
    #     "phone": "050-5894-1192",
    #     "website": "https://www.makeyousmile.jp/",
    #     "rep_email": "kazuhiro110@makeyousmile.jp",
    #     "agency_name": "AQIUM INTERNATIONAL INC.",
    #     "agency_address": "Unit 2A, 3A & 3B, 4K Plaza Bldg., 677 Shaw Blvd., Brgy. Kapitolyo, Pasig City, Philippines",
    #     "agency_rep_name": "Ms. HERMARANOEMI B. ALBERTO",
    #     "agency_rep_position": "Director",
    #     "visa_type": "Engineer / Specialist in Humanities / International Services",
    # }

    replacements = {
        "created_date": f"{day} {month_name}, {year}",
        "company_name_en": company.name_en,
        "rep_name_en": company.rep_name_en,
        "rep_position_en": company.rep_position_en,
        "building": company.building_en,
        # "street": "4-20-2 Kaigandori",
        "street": "",
        "city": company.municipality_town_en,
        "prefecture": company.prefecture_en,
        "postal_code": company.postal_code,
        "phone": company.phone,
        "website": company.website,
        "rep_email": company.rep_email,
        "agency_name": agency.name,
        "agency_address": agency.address,
        "agency_rep_name": agency.rep_name,
        "agency_rep_position": agency.rep_position,
        "visa_type": details['visa_type'],
    }

    job_details = details['job_details']
    # total_workers = sum(detail['no_of_workers'] for detail in job_details)
    total_workers = details['total_workers']

    with TemporaryDirectory() as temp_dir:
        temp_file_path = os.path.join(temp_dir, new_document_name)

        new_contract_buffer = BytesIO()
        fill_manpower_request_template(
            original_file, temp_file_path, replacements, job_details, total_workers)

        with open(temp_file_path, 'rb') as temp_file:
            new_contract_buffer.write(temp_file.read())

        s3.upload_file(temp_file_path, Bucket=bucket_name, Key=s3_new_document, ExtraArgs={
            'ACL': 'public-read', 'ContentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})

    s3_read_url = generate_s3_url(s3_new_document, 'read')
    return s3_read_url


def fill_employment_contract(company, agency, details):
    # created_at = datetime.now()
    # year = created_at.year
    # month = created_at.month
    # day = created_at.day

    if details['created_date'] is None or details['created_date'] == "":
        created_date = datetime.now()
    else:
        created_date = datetime.strptime(
            details['created_date'], "%Y-%m-%dT%H:%M:%S.%fZ")

    year = created_date.year
    month = created_date.month
    month_name = created_date.strftime("%b")
    day = created_date.day

    # current date
    today = datetime.now()

    # new_document_name = "employment_contract_test.docx"
    new_document_name = f"MWO_EMPLOYMENT_CONTRACT_{company.name_en}_{today.strftime('%Y%m%d')}.docx"
    original_file = os.path.join(STATIC_DIR, "employment_contract.docx")
    s3_new_document = f"{contracts_folder}{new_document_name}"

    # replacements = {
    #     "created_date": f"{day} {month}, {year}",
    #     "company_name_en": "Make You Smile Co., Ltd.",
    #     "company_address": "#304 YT Bashamichi Bldg, 4-20-2 Kaigandori, Naka-ku, Yokohama shi, Kanagawa Pref.",
    #     "company_rep_name_en": "Kazuhiro Ito",
    #     "company_rep_position_en": "President",
    #     "building": "#304 YT Bashamichi Bldg",
    #     "street": "4-20-2 Kaigandori",
    #     "city": "Naka-ku, Yokohama shi",
    #     "prefecture": "Kanagawa Pref.",
    #     "postal_code": "231-0002",
    #     "company_phone": "050-5894-1192",
    #     "website": "https://www.makeyousmile.jp/",
    #     "rep_email": "kazuhiro110makeyousmile.jp",
    #     "agency_name": "AQIUM INTERNATIONAL INC.",
    #     "agency_address": "Unit 2A, 3A & 3B, 4K Plaza Bldg., 677 Shaw Blvd., Brgy. Kapitolyo, Pasig City, Philippines",
    #     "agency_rep_name": "Ms. HERMARANOEMI B. ALBERTO",
    #     "agency_rep_position": "Director",
    #     "employment_address": "Japan",
    #     "employment_term": "3 years",
    #     "job_position_title": "Software Engineer",
    #     "job_position_description": "Develop software applications",
    # }

    if details['passport_date_issued'] is None or details['passport_date_issued'] == "":
        pdi_year = ""
        pdi_month = ""
        pdi_month_name = ""
        pdi_day = ""
    else:
        passport_date_issued = datetime.strptime(
            details['passport_date_issued'], "%Y-%m-%dT%H:%M:%S.%fZ")

        pdi_year = passport_date_issued.year
        pdi_month = passport_date_issued.month
        pdi_month_name = passport_date_issued.strftime("%b")
        pdi_day = passport_date_issued.day

    replacements = {
        "created_date": f"{day} {month_name}, {year}",
        "company_name_en": company.name_en,
        "company_address": f"{company.building_en}, {company.municipality_town_en}, {company.prefecture_en}",
        "company_rep_name_en": company.rep_name_en,
        "company_rep_position_en": company.rep_position_en,
        "building": company.building_en,
        "street": "",
        "city": company.municipality_town_en,
        "prefecture": company.prefecture_en,
        "postal_code": company.postal_code,
        "company_phone": company.phone,
        "website": company.website,
        "rep_email": company.rep_email,
        "agency_name": agency.name,
        "agency_address": agency.address,
        "agency_rep_name": agency.rep_name,
        "agency_rep_position": agency.rep_position,
        "worker_name": details['worker_name'],
        "philippine_address": details['philippine_address'],
        "civil_status": details['civil_status'],
        "passport_no": details['passport_no'],
        "passport_date_issued": f"{pdi_month_name} {pdi_day}, {pdi_year}",
        "passport_place_issued": details['passport_place_issued'],
        "employment_address": details['employment_address'],
        "employment_term": details['employment_term'],
        "job_position_title": details['job_position_title'],
        "job_position_description": details['job_position_description'],
    }

    with TemporaryDirectory() as temp_dir:
        temp_file_path = os.path.join(temp_dir, new_document_name)

        new_contract_buffer = BytesIO()
        fill_template(original_file, temp_file_path, replacements)

        with open(temp_file_path, 'rb') as temp_file:
            new_contract_buffer.write(temp_file.read())

        s3.upload_file(temp_file_path, Bucket=bucket_name, Key=s3_new_document, ExtraArgs={
            'ACL': 'public-read', 'ContentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})

    s3_read_url = generate_s3_url(s3_new_document, 'read')
    return s3_read_url


def fill_recruitment_agreement(company, agency, details):
    if details['created_date'] is None or details['created_date'] == "":
        created_date = datetime.now()
    else:
        created_date = datetime.strptime(
            details['created_date'], "%Y-%m-%dT%H:%M:%S.%fZ")

    year = created_date.year
    month = created_date.month
    month_name = created_date.strftime("%B")
    day = created_date.day

    # current date
    today = datetime.now()

    visa_type = details['visa_type']

    if visa_type == 'psw':
        # new_document_name = f"MWO_雇用契約書_{company.name_en}_{created_date.strftime('%Y%m%d')}.docx"
        new_document_name = f"MWO_RECRUITMENTAGREEMENT_{company.name_en}_{today.strftime('%Y%m%d')}.docx"
        original_file = os.path.join(STATIC_DIR, "recruitment_agreement.docx")
    elif visa_type == 'ssw':
        # new_document_name = f"MWO_雇用契約書_{company.name_en}_{created_date.strftime('%Y%m%d')}.docx"
        new_document_name = f"MWO_SSW_RECRUITMENTAGREEMENT_{company.name_en}_{today.strftime('%Y%m%d')}.docx"
        original_file = os.path.join(
            STATIC_DIR, "ssw", "recruitment_agreement.docx")

    s3_new_document = f"{contracts_folder}{new_document_name}"

    replacements = {
        "created_date": f"{day} {month_name}, {year}",
        "year": year,
        "month": month,
        "month_name": month_name,
        "day": day,
        "company_name_en": company.name_en,
        "company_address_en": f"{company.building_en}, {company.municipality_town_en}, {company.prefecture_en}",
        "company_rep_name_en": company.rep_name_en,
        "company_rep_position_en": company.rep_position_en,
        "company_name_ja": company.name_ja,
        "company_address_ja": f"{company.prefecture_ja}, {company.municipality_town_ja} {company.street_address_ja}, {company.building_ja}",
        "company_rep_name_ja": company.rep_name_ja,
        "company_rep_position_ja": company.rep_position_ja,
        "postal_code": company.postal_code,
        "rep_email": company.rep_email,
        "agency_name": agency.name,
        "agency_address": agency.address,
        "agency_rep_name": agency.rep_name,
        "agency_rep_position": agency.rep_position,
    }

    with TemporaryDirectory() as temp_dir:
        temp_file_path = os.path.join(temp_dir, new_document_name)

        new_contract_buffer = BytesIO()
        fill_template(original_file, temp_file_path, replacements)

        with open(temp_file_path, 'rb') as temp_file:
            new_contract_buffer.write(temp_file.read())

        s3.upload_file(temp_file_path, Bucket=bucket_name, Key=s3_new_document, ExtraArgs={
            'ACL': 'public-read', 'ContentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})

    s3_read_url = generate_s3_url(s3_new_document, 'read')
    return s3_read_url

# ================================ SSW ================================


def fill_ssw_company_profile(company, agency, details):
    if details['created_date'] is None or details['created_date'] == "":
        created_date = datetime.now()
    else:
        created_date = datetime.strptime(
            details['created_date'], "%Y-%m-%dT%H:%M:%S.%fZ")

    year = created_date.year
    month = created_date.month
    month_name = created_date.strftime("%B")
    day = created_date.day

    # current date
    today = datetime.now()

    # new_document_name = f"MWO_雇用契約書_{company.name_en}_{created_date.strftime('%Y%m%d')}.docx"
    new_document_name = f"MWO_SSW_COMPANYPROFILE_{company.name_en}_{today.strftime('%Y%m%d')}.docx"
    original_file = os.path.join(STATIC_DIR, "ssw", "company_profile.docx")
    s3_new_document = f"{contracts_folder}{new_document_name}"

    replacements = {
        "created_date": f"{day} {month_name}, {year}",
        "year": year,
        "month": month,
        "month_name": month_name,
        "day": day,
        "company_name_en": company.name_en,
        "company_address_en": f"{company.building_en}, {company.municipality_town_en}, {company.prefecture_en}",
        "rep_name_en": company.rep_name_en,
        "rep_position_en": company.rep_position_en,
        "company_name_ja": company.name_ja,
        "company_address_ja": f"{company.prefecture_ja}, {company.municipality_town_ja} {company.street_address_ja}, {company.building_ja}",
        "company_rep_name_ja": company.rep_name_ja,
        "company_rep_position_ja": company.rep_position_ja,
        "postal_code": company.postal_code,
        "address": f"{company.building_en}, {company.municipality_town_en}, {company.prefecture_en}",
        "phone": company.phone,
        "email": company.email,
        "website": company.website,
        "rep_email": company.rep_email,
        "agency_name": agency.name,
        "agency_address": agency.address,
        "agency_rep_name": agency.rep_name,
        "agency_rep_position": agency.rep_position,
        "year_established": company.year_established if company.year_established is not None else "",
        "registered_industry": company.registered_industry_en if company.registered_industry_en is not None else "",
        "regular_worker_count": str(company.regular_worker_count) if company.regular_worker_count is not None else "",
        "parttime_worker_count": str(company.parttime_worker_count) if company.parttime_worker_count is not None else "",
        "foreigner_worker_count": str(company.foreigner_worker_count) if company.foreigner_worker_count is not None else "",
    }

    with TemporaryDirectory() as temp_dir:
        temp_file_path = os.path.join(temp_dir, new_document_name)

        new_contract_buffer = BytesIO()
        fill_template(original_file, temp_file_path, replacements)

        with open(temp_file_path, 'rb') as temp_file:
            new_contract_buffer.write(temp_file.read())

        s3.upload_file(temp_file_path, Bucket=bucket_name, Key=s3_new_document, ExtraArgs={
            'ACL': 'public-read', 'ContentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})

    s3_read_url = generate_s3_url(s3_new_document, 'read')
    return s3_read_url


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """
    Set margins for a cell in a table. Margins are set in TWIPs.
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Create <w:tcMar> element if it doesn't exist
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)

    # Set the margin values
    for margin_type, margin_value in [('top', top), ('left', start), ('bottom', bottom), ('right', end)]:
        margin_tag = f'w:{margin_type}'
        margin_element = tcMar.find(qn(margin_tag))
        if margin_element is None:
            margin_element = OxmlElement(margin_tag)
            tcMar.append(margin_element)
        margin_element.set(qn('w:w'), str(margin_value))
        margin_element.set(qn('w:type'), 'dxa')


def fill_task_list_template(template_path, output_path, replacements, task_duties_details):
    doc = Document(template_path)

    # Replace placeholders in the document
    for placeholder, replacement in replacements.items():
        replace_placeholder(doc, f"{{{{{placeholder}}}}}", replacement)

    # Assuming the table is where you want to add the tasks and qualifications
    task_table = doc.tables[0]

    # Add header row for "A. List of Tasks/Assignment"
    header_row_a = task_table.add_row()
    # Merge cells for the header
    header_row_a.cells[0].merge(header_row_a.cells[1])
    header_row_a.cells[0].text = "A. List of Tasks/Assignment to be performed by the Filipino SSW"
    set_cell_margins(header_row_a.cells[0],
                     top=100, start=100, bottom=100, end=100)

    # Insert the task list under the header
    tasks = task_duties_details["tasksList"]
    for i, task in enumerate(tasks, start=1):
        new_row = task_table.add_row()
        new_row.cells[0].text = f"{i}."
        new_row.cells[1].text = task

        # Set margins for cells
        set_cell_margins(new_row.cells[0], top=100,
                         start=100, bottom=100, end=100)
        set_cell_margins(new_row.cells[1], top=100,
                         start=100, bottom=100, end=100)

    # Add header row for "B. List of Additional Criteria/Qualifications"
    header_row_b = task_table.add_row()
    # Merge cells for the header
    header_row_b.cells[0].merge(header_row_b.cells[1])
    header_row_b.cells[0].text = "B. List of Additional Criteria/Qualifications"
    set_cell_margins(header_row_b.cells[0],
                     top=100, start=100, bottom=100, end=100)

    # Insert the qualifications list under the header
    qualifications = task_duties_details["qualificationsList"]
    for i, qual in enumerate(qualifications, start=1):
        new_row = task_table.add_row()
        new_row.cells[0].text = f"{i}."
        new_row.cells[1].text = qual

        # Set margins for cells
        set_cell_margins(new_row.cells[0], top=100,
                         start=100, bottom=100, end=100)
        set_cell_margins(new_row.cells[1], top=100,
                         start=100, bottom=100, end=100)

    # Save the document
    doc.save(output_path)


def fill_ssw_list_task_duties(company, agency, details):
    if details['created_date'] is None or details['created_date'] == "":
        created_date = datetime.now()
    else:
        created_date = datetime.strptime(
            details['created_date'], "%Y-%m-%dT%H:%M:%S.%fZ")

    year = created_date.year
    month = created_date.month
    month_name = created_date.strftime("%B")
    day = created_date.day

    # Current date
    today = datetime.now()

    # Generate the new document name
    new_document_name = f"MWO_SSW_LISTTASKSQUALIFICATIONS_{company.name_en}_{today.strftime('%Y%m%d')}.docx"
    original_file = os.path.join(STATIC_DIR, "ssw", "task_criteria_list.docx")
    s3_new_document = f"{contracts_folder}{new_document_name}"

    tasksDutiesListObjects = [
        {
            "id": 1,
            "job_code": "care_worker",
            "job_title": "CARE WORKER",
            "job_description": "Assist special needs patients with special needs with their needs. Caregivers help patients take their medications, bathe, dress, eat, groom, and take care of their personal needs according to the care plan set forth in the group home.",
            "tasksList": ["Assisting clients with care such as grooming, washing, and maintaining personal hygiene standards.", "Assess medical needs. Monitoring Individuals conditions by taking their temperature, pulse, respiration and weight, and possibly helping medication.", "Following a prescribed healthcare plan, which may include assisting with exercise and administering medications.", "Catheter care, Bowel Care, Use of eye drops, Application of ointments; Health-related tasks may be required at any time of night or day.", "Domestic Cleaning (including vacuuming, dusting, polishing, cleaning bathrooms, kitchens and etc.) Maintaining to a clean, well-organized and pleasant environment for all patients.", "Preparing food and drinks for patients. Preparing and serving healthy, nutritious meals, taking into account any special dietary needs and personal preferences.", "Escorting on activities outside the home. Providing mobility assistance maybe required, for example helping the patience in and out of bed, a chair or a wheelchair.", "Providing emotional support and encouragement to perform necessary tasks.", "Specialist Support; People with mental Heath needs, Dementia, Learning Disabilities, Physical Disabilities, People who are terminally ill.", "Care Worker may be asked to carry out any other duties that may reasonably be required by management. Ensuring that you meet all the required standards of health and safety.", "Must Protect the Privacy of the clients. And Making sure that your client is comfortable, happy and well cared for all the times."],
            "qualificationsList": ["An applicant must be at least a High School Diploma. And Further education in Caregiving NC2 is very beneficial", "23 years and above but not more than 45 years old", "Japanese Language Proficiency test (N4 or Higher), and have the ability to communicate BASIC NIHONGO and Daily Conversation", "Successfully finished the 3 years as a TRAINEE and pass the grade 3 Skill exam or SSW Caregiver PROMETRIC EXAM both English and Japanese", "Must have passed the SSW Care Worker Prometric Exam in English and Japanese.", "A professional and friendly attitude, good sense of Humor and approachable demeanor. Have a sense of calm, especially when under pressure, as some clients can be anxious, nervous and even aggressive", "The ability to build strong relationship with clients and other families, and build a sense of Trust nearly limitless patience.", "A willingness to work flexible hours, which may include night shift.", "Sensitivity towards potentially embarrassing medical conditions", "Can respect personal living habits, culture and sense of values."]
        },
        {
            "id": 2,
            "job_code": "restaurant_staff",
            "job_title": "RESTAURANT STAFF",
            "job_description": "Responsible for taking orders from patrons and delivering their food to them. Dealing with customers and kitchen staffs and cleaning and sanitizing the workspace.",
            "tasksList": ["Welcome customers and lead them to their seats.", "Introduce restaurant menu.", "Deliver orders to the table in timely manner. Inspect the quality and looks of dishes, making sure they uphold standards.", "Cleaning the work space or workstation and food storage areas.", "Assist in food preparation."],
            "qualificationsList": ["Ability to follow routine verbal and written instructions in Japanese language", "Japanese Language Proficiency Test at least N4 level passer or Prometric Japan Foundation Test-Basic passer", "Prometric Food service skill test passer (Japanese)", "Physically fit and healthy", "At least 6 months Food service industry experience"]
        },
        {
            "id": 3,
            "job_code": "building_cleaner",
            "job_title": "INTERIOR BUILDING CLEANER",
            "job_description": "Keeping the building or establishment clean and organized.",
            "tasksList": ["Responsible for all basic cleaning in and around the building premises ", "Sets up stocks and maintain cleaning equipment", "Monitors and maintain sanitation and organization of assigned areas "],
            "qualificationsList": ["College / High School / Vocational Graduate", "Japan Foundation Test (JFT-Basic) Passer / JLPT N4 or above", "Prometric Building Cleaning Management Proficiency Test Passer / Former TITP in the same category"]
        }
    ]

    selected_position = details["ssw_job_title"]
    selected_task_duties = next(
        (item for item in tasksDutiesListObjects if item["job_code"] == selected_position), None)

    # Update replacements based on the selected job title and description
    replacements = {
        "company_name_en": company.name_en,
        "company_address_en": f"{company.building_en}, {company.municipality_town_en}, {company.prefecture_en}",
        "rep_name_en": company.rep_name_en,
        "rep_position_en": company.rep_position_en,
        "job_position": selected_task_duties["job_title"] if selected_task_duties else "",
        "job_description": selected_task_duties["job_description"] if selected_task_duties else ""
    }

    with TemporaryDirectory() as temp_dir:
        temp_file_path = os.path.join(temp_dir, new_document_name)

        new_contract_buffer = BytesIO()
        fill_task_list_template(
            original_file, temp_file_path, replacements, selected_task_duties)

        with open(temp_file_path, 'rb') as temp_file:
            new_contract_buffer.write(temp_file.read())

        s3.upload_file(temp_file_path, Bucket=bucket_name, Key=s3_new_document, ExtraArgs={
            'ACL': 'public-read', 'ContentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})

    s3_read_url = generate_s3_url(s3_new_document, 'read')
    return s3_read_url

# ================================ SSW ================================


def generate_default_documents(document_name: str, application_type: str):
    if application_type == "initial":
        if document_name == 'aqium_license_copy':
            # retrieve the aqium_license_copy.pdf file, upload to s3 and return the s3 url
            new_document_name = "★AQIUMライセンスの写し.pdf"
            original_file = os.path.join(STATIC_DIR, "aqium_license_copy.pdf")

            s3_new_document = f"{contracts_folder}{new_document_name}"

            with open(original_file, 'rb') as file:
                s3.upload_fileobj(file, bucket_name, s3_new_document, ExtraArgs={
                    'ACL': 'public-read', 'ContentType': 'application/pdf'})

            s3_read_url = generate_s3_url(s3_new_document, 'read')

            return s3_read_url
        elif document_name == 'aqium_representative_passport_copy':
            # retrieve the aqium_representative_passport_copy.pdf file, upload to s3 and return the s3 url
            new_document_name = "★AQIUM代表者のパスポートの写し.pdf"
            original_file = os.path.join(STATIC_DIR, "noemi_passport_copy.pdf")

            s3_new_document = f"{contracts_folder}{new_document_name}"

            with open(original_file, 'rb') as file:
                s3.upload_fileobj(file, bucket_name, s3_new_document, ExtraArgs={
                    'ACL': 'public-read', 'ContentType': 'application/pdf'})

            s3_read_url = generate_s3_url(s3_new_document, 'read')

            return s3_read_url
        elif document_name == 'psw_initial_checklist':
            # retrieve the aqium_representative_passport_copy.pdf file, upload to s3 and return the s3 url
            new_document_name = "新規「技術・人文知識・国際業務」要件のチェックリスト ガイド.docx"
            original_file = os.path.join(
                STATIC_DIR, "psw_initial_checklist.docx")

            s3_new_document = f"{contracts_folder}{new_document_name}"

            # upload file to s3
            with open(original_file, 'rb') as file:
                s3.upload_fileobj(file, bucket_name, s3_new_document, ExtraArgs={
                    'ACL': 'public-read', 'ContentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})

            s3_read_url = generate_s3_url(s3_new_document, 'read')

            return s3_read_url
        elif document_name == 'ssw_initial_checklist':
            # retrieve the aqium_representative_passport_copy.pdf file, upload to s3 and return the s3 url
            new_document_name = "新規「特定技能」要件のチェックリスト ガイド.docx"
            original_file = os.path.join(
                STATIC_DIR, "ssw_initial_checklist.docx")

            s3_new_document = f"{contracts_folder}{new_document_name}"

            with open(original_file, 'rb') as file:
                # upload using upload_file, this is not pdf, it is docx
                s3.upload_fileobj(file, bucket_name, s3_new_document, ExtraArgs={
                    'ACL': 'public-read', 'ContentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})

            s3_read_url = generate_s3_url(s3_new_document, 'read')

            return s3_read_url

    else:
        raise ValueError("File not found")


# def generate_letter_pack(details, sheet_configs):
#     # Load the workbook
#     document_path = os.path.join(STATIC_DIR, "letter_pack.xlsx")
#     wb = openpyxl.load_workbook(document_path, data_only=False)  # Ensure formulas are not replaced by their values

#     recipient_id = details['recipient_id']

#     for sheet_name, sheet_details in sheet_configs.items():
#         ws = wb[sheet_name]

#         ws['V4'] = str(recipient_id)

#         if '日本語' in sheet_name:
#             ws['G14'] = sheet_details['sender_address_key']
#             ws['G16'] = sheet_details['sender_name_key']
#             ws['G17'] = sheet_details['sender_rep_name_key']

#             if "-" not in sheet_details['sender_tel_key']:
#                 ws.merge_cells('H20:L20')
#                 ws['H20'] = sheet_details['sender_tel_key']
#             else:
#                 try:
#                     ws.unmerge_cells('H20:L20')
#                 except:  # noqa: E722
#                     pass  # continue if not merged

#                 tel_parts = sheet_details['sender_tel_key'].split('-')
#                 ws['H20'] = tel_parts[0]
#                 ws['J20'] = tel_parts[1]
#                 ws['L20'] = tel_parts[2]

#         elif "英語" in sheet_name:
#             ws['G14'] = sheet_details['sender_address_key']
#             ws['G15'] = sheet_details['sender_address_key2']
#             ws['G17'] = sheet_details['sender_name_key']
#             ws['G18'] = sheet_details['sender_rep_name_key']

#             if "-" not in sheet_details['sender_tel_key']:
#                 ws.merge_cells('H20:P20')
#                 ws['H20'] = sheet_details['sender_tel_key']
#             else:
#                 try:
#                     ws.unmerge_cells('H20:P20')
#                 except:  # noqa: E722
#                     pass  # continue if not merged

#                 tel_parts = sheet_details['sender_tel_key'].split('-')
#                 ws['H20'] = tel_parts[0]
#                 ws['K20'] = tel_parts[1]
#                 ws['P20'] = tel_parts[2]

#     # Save the modified workbook back to the original file
#     wb.save(document_path)

#     # Upload the modified workbook directly to S3
#     s3 = boto3.client('s3')
#     with open(document_path, 'rb') as file:
#         s3.upload_fileobj(file, bucket_name, f"{contracts_folder}letter_pack.xlsx", ExtraArgs={
#             'ACL': 'public-read', 'ContentType': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'})

#     s3_read_url = generate_s3_url(f"{contracts_folder}letter_pack.xlsx", 'read')

#     return s3_read_url

def generate_letter_pack(details, sheet_configs):
    document_path = os.path.join(STATIC_DIR, "letter_pack.xlsx")
    app = xw.App(visible=False)
    wb = app.books.open(document_path)

    recipient_id = details['recipient_id']

    for sheet_name, sheet_details in sheet_configs.items():
        ws = wb.sheets[sheet_name]

        ws.range('V4').value = str(recipient_id)

        if '日本語' in sheet_name:
            ws.range('G14').value = sheet_details['sender_address_key']
            ws.range('G16').value = sheet_details['sender_name_key']
            ws.range('G17').value = sheet_details['sender_rep_name_key']

            if "-" not in sheet_details['sender_tel_key']:
                ws.range('H20:L20').merge()
                ws.range('H20').value = sheet_details['sender_tel_key']
            else:
                ws.range('H20:L20').unmerge()
                tel_parts = sheet_details['sender_tel_key'].split('-')
                ws.range('H20').value = tel_parts[0]
                ws.range('J20').value = tel_parts[1]
                ws.range('L20').value = tel_parts[2]

        elif "英語" in sheet_name:
            ws.range('G14').value = sheet_details['sender_address_key']
            ws.range('G15').value = sheet_details['sender_address_key2']
            ws.range('G17').value = sheet_details['sender_name_key']
            ws.range('G18').value = sheet_details['sender_rep_name_key']

            if "-" not in sheet_details['sender_tel_key']:
                ws.range('H20:P20').merge()
                ws.range('H20').value = sheet_details['sender_tel_key']
            else:
                ws.range('H20:P20').unmerge()
                tel_parts = sheet_details['sender_tel_key'].split('-')
                ws.range('H20').value = tel_parts[0]
                ws.range('K20').value = tel_parts[1]
                ws.range('P20').value = tel_parts[2]

    # # Ensure all shapes and text boxes are preserved
    # for sheet in wb.sheets:
    #     shapes = sheet.shapes
    #     for shape in shapes:
    #         print(f"Shape '{shape.name}' found on sheet '{sheet.name}'")

    # Save and close the workbook
    wb.save(document_path)
    wb.close()
    app.quit()

    # Upload the modified workbook directly to S3
    s3 = boto3.client('s3')
    with open(document_path, 'rb') as file:
        s3.upload_fileobj(file, bucket_name, f"{contracts_folder}レターパック宛先.xlsx", ExtraArgs={
            'ACL': 'public-read', 'ContentType': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'})

    s3_read_url = generate_s3_url(f"{contracts_folder}レターパック宛先.xlsx", 'read')

    return s3_read_url